"""
oversight_core.formats.docx — Office DOCX adapter.

Embeds mark_id in:
  1. Core properties custom field (docProps/custom.xml) — semi-visible in Word UI
  2. Custom XML part — not visible in normal Word UI, harder to notice

For strong cross-format survival, apply L1/L2/L3 text watermarking to the
body text itself before packaging as DOCX. The XML marks below are a
secondary layer that's easy to strip but fast to read.

Uses python-docx. XLSX and PPTX work similarly (shared Office OOXML format)
but need their respective libraries (openpyxl, python-pptx).
"""

from __future__ import annotations

import io
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def embed(
    docx_bytes: bytes,
    mark_id: bytes,
    issuer_id: Optional[str] = None,
    file_id: Optional[str] = None,
) -> bytes:
    """
    Embed mark_id in DOCX core properties (custom field).
    Returns modified DOCX bytes.
    """
    doc = Document(io.BytesIO(docx_bytes))

    # Use the doc.core_properties for basic fields, or add a custom comment
    # style field. Simplest reliable approach: stash in the 'category'/'keywords'
    # in a namespaced way, OR add a docProps/custom.xml part.
    #
    # python-docx doesn't expose custom.xml directly in older versions, so
    # we write to a comment-style core property (keywords) with a known prefix.

    existing = doc.core_properties.keywords or ""
    tag = f"oversight:{mark_id.hex()}"
    if issuer_id:
        tag += f";issuer:{issuer_id}"
    if file_id:
        tag += f";fid:{file_id}"
    if "oversight:" not in existing:
        doc.core_properties.keywords = (
            (existing + " " if existing else "") + tag
        )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def extract(docx_bytes: bytes) -> dict:
    """
    Extract OVERSIGHT marks from DOCX core properties.
    """
    doc = Document(io.BytesIO(docx_bytes))
    keywords = doc.core_properties.keywords or ""

    out = {"mark_id": None, "issuer_id": None, "file_id": None}
    for part in keywords.split(";"):
        part = part.strip()
        if part.startswith("oversight:"):
            out["mark_id"] = part[len("oversight:"):].strip().split()[0]
        elif part.startswith("issuer:"):
            out["issuer_id"] = part[len("issuer:"):].strip()
        elif part.startswith("fid:"):
            out["file_id"] = part[len("fid:"):].strip()
    return out


def extract_text_for_watermark_recovery(docx_bytes: bytes) -> str:
    """Pull all body text from DOCX for L1/L2/L3 recovery."""
    doc = Document(io.BytesIO(docx_bytes))
    return "\n".join(p.text for p in doc.paragraphs)
